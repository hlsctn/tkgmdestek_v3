import sys
import sqlite3
import os
from PyQt6.QtWidgets import (
    QApplication, QDialog, QLabel, QVBoxLayout, QLineEdit, QTextEdit,
    QListWidget, QPushButton, QHBoxLayout, QFileDialog
)
from PyQt6.QtGui import QFont
from PyQt6.QtCore import Qt
from docx import Document

# ---------------------------
# Veritabanı yolu ve klasör kontrolü
# ---------------------------
PROJE_KLASOR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
VERITABANI_KLASOR = os.path.join(PROJE_KLASOR, "veritabani")
os.makedirs(VERITABANI_KLASOR, exist_ok=True)
DB_YOLU = os.path.join(VERITABANI_KLASOR, "notlar.db")

# ---------------------------
# Veritabanı Fonksiyonları
# ---------------------------
def veritabani_olustur():
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS notlar (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            baslik TEXT NOT NULL,
            icerik TEXT
        )
    """)
    conn.commit()
    conn.close()

def notlari_getir():
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("SELECT id, baslik, icerik FROM notlar")
    veriler = cursor.fetchall()
    conn.close()
    return veriler

def not_ekle_db(baslik, icerik):
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO notlar (baslik, icerik) VALUES (?, ?)", (baslik, icerik))
    conn.commit()
    conn.close()

def not_guncelle_db(not_id, baslik, icerik):
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("UPDATE notlar SET baslik=?, icerik=? WHERE id=?", (baslik, icerik, not_id))
    conn.commit()
    conn.close()

def not_sil_db(not_id):
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM notlar WHERE id=?", (not_id,))
    conn.commit()
    conn.close()

# ---------------------------
# Not Defteri Penceresi
# ---------------------------
class NotlarimPenceresi(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Notlarım")
        self.setGeometry(0, 0, 550, 500)
        self.center()
        self.secili_not_id = None

        # Ana layout
        layout = QVBoxLayout()
        self.setLayout(layout)
        self.setStyleSheet("background-color: #f2f2f2; color: black;")

        # Başlık label
        baslikLabel = QLabel("Notlarım")
        baslikLabel.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        baslikLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(baslikLabel)

        # Başlık input
        self.baslikInput = QLineEdit()
        self.baslikInput.setPlaceholderText("Başlık girin...")
        self.baslikInput.setStyleSheet(
            "padding: 8px; border-radius: 5px; border: 1px solid #ccc; color: black;"
        )
        layout.addWidget(self.baslikInput)

        # Detay input
        self.detayInput = QTextEdit()
        self.detayInput.setPlaceholderText("Detayları buraya yazın...")
        self.detayInput.setStyleSheet(
            "padding: 8px; border-radius: 5px; border: 1px solid #ccc; color: black;"
        )
        layout.addWidget(self.detayInput)

        # Not listesi
        self.notList = QListWidget()
        self.notList.setStyleSheet("padding:5px; color: black;")
        self.notList.itemClicked.connect(self.not_secili)
        layout.addWidget(self.notList)

        # Butonlar
        btnLayout = QHBoxLayout()
        layout.addLayout(btnLayout)

        self.btnKaydet = QPushButton("Kaydet / Güncelle")
        self.btnKaydet.setStyleSheet(self.button_style())
        self.btnKaydet.clicked.connect(self.not_kaydet)
        btnLayout.addWidget(self.btnKaydet)

        self.btnSil = QPushButton("Seçili Notu Sil")
        self.btnSil.setStyleSheet(self.button_style())
        self.btnSil.clicked.connect(self.not_sil)
        btnLayout.addWidget(self.btnSil)

        self.btnIndir = QPushButton("Notu İndir (.docx)")
        self.btnIndir.setStyleSheet(self.button_style())
        self.btnIndir.clicked.connect(self.not_indir)
        btnLayout.addWidget(self.btnIndir)

        # Veritabanından notları yükle
        self.notlari_yukle()

    # Modern buton stili
    def button_style(self):
        return (
            "QPushButton {background-color: #ff9800; color: white; padding: 8px; border-radius: 5px;}"
            "QPushButton:hover {background-color: #e68900;}"
        )

    # Pencereyi ortala
    def center(self):
        screen = self.screen()
        rect = screen.availableGeometry()
        x = rect.x() + (rect.width() - self.width()) // 2
        y = rect.y() + (rect.height() - self.height()) // 2
        self.move(x, y)

    # Notları yükle
    def notlari_yukle(self):
        self.notList.clear()
        for not_id, baslik, icerik in notlari_getir():
            self.notList.addItem(f"{not_id} - {baslik}")

    # Liste öğesi seçildi
    def not_secili(self, item):
        not_id = int(item.text().split(" - ")[0])
        conn = sqlite3.connect(DB_YOLU)
        cursor = conn.cursor()
        cursor.execute("SELECT baslik, icerik FROM notlar WHERE id=?", (not_id,))
        baslik, icerik = cursor.fetchone()
        conn.close()
        self.baslikInput.setText(baslik)
        self.detayInput.setPlainText(icerik)
        self.secili_not_id = not_id

    # Kaydet / Güncelle
    def not_kaydet(self):
        baslik = self.baslikInput.text().strip()
        icerik = self.detayInput.toPlainText().strip()
        if not baslik:
            return
        if self.secili_not_id:
            not_guncelle_db(self.secili_not_id, baslik, icerik)
        else:
            not_ekle_db(baslik, icerik)
        self.baslikInput.clear()
        self.detayInput.clear()
        self.secili_not_id = None
        self.notlari_yukle()

    # Sil
    def not_sil(self):
        if self.secili_not_id:
            not_sil_db(self.secili_not_id)
            self.baslikInput.clear()
            self.detayInput.clear()
            self.secili_not_id = None
            self.notlari_yukle()

    # DOCX olarak kaydet (Masaüstü varsayılan)
    def not_indir(self):
        if not self.secili_not_id:
            return
        baslik = self.baslikInput.text()
        icerik = self.detayInput.toPlainText()
        options = QFileDialog.Option.DontUseNativeDialog  # PyQt6 uyumlu
        masaustu = os.path.join(os.path.expanduser("~"), "Desktop")
        dosya_yolu, _ = QFileDialog.getSaveFileName(
            self,
            "Notu Kaydet",
            os.path.join(masaustu, f"{baslik}.docx"),  # masaüstü varsayılan
            "Word Documents (*.docx)",
            options=options
        )
        if dosya_yolu:
            if not dosya_yolu.lower().endswith(".docx"):
                dosya_yolu += ".docx"
            doc = Document()
            doc.add_heading(baslik, level=1)
            doc.add_paragraph(icerik)
            doc.save(dosya_yolu)

# ---------------------------
# Programı çalıştır
# ---------------------------
if __name__ == "__main__":
    veritabani_olustur()
    app = QApplication(sys.argv)
    pencere = NotlarimPenceresi()
    pencere.show()
    sys.exit(app.exec())
