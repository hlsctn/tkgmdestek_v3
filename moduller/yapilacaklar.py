import sys
import sqlite3
import os
from PyQt6.QtWidgets import (
    QApplication, QDialog, QLabel, QVBoxLayout, QLineEdit, QTextEdit,
    QListWidget, QListWidgetItem, QPushButton, QHBoxLayout, QFileDialog, QComboBox
)
from PyQt6.QtGui import QFont, QColor, QBrush
from PyQt6.QtCore import Qt
from docx import Document

# ---------------------------
# Veritabanı yolu ve klasör kontrolü
# ---------------------------
PROJE_KLASOR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
VERITABANI_KLASOR = os.path.join(PROJE_KLASOR, "veritabani")
os.makedirs(VERITABANI_KLASOR, exist_ok=True)
DB_YOLU = os.path.join(VERITABANI_KLASOR, "yapilacaklar.db")

# ---------------------------
# Veritabanı Fonksiyonları
# ---------------------------
def veritabani_olustur():
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("DROP TABLE IF EXISTS yapilacaklar")
    cursor.execute("""
        CREATE TABLE yapilacaklar (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            baslik TEXT NOT NULL,
            icerik TEXT,
            onem INTEGER DEFAULT 2
        )
    """)
    conn.commit()
    conn.close()

def gorevleri_getir():
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("SELECT id, baslik, icerik, onem FROM yapilacaklar")
    veriler = cursor.fetchall()
    conn.close()
    return veriler

def gorev_ekle_db(baslik, icerik, onem):
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO yapilacaklar (baslik, icerik, onem) VALUES (?, ?, ?)", (baslik, icerik, onem))
    conn.commit()
    conn.close()

def gorev_guncelle_db(gorev_id, baslik, icerik, onem):
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("UPDATE yapilacaklar SET baslik=?, icerik=?, onem=? WHERE id=?", (baslik, icerik, onem, gorev_id))
    conn.commit()
    conn.close()

def gorev_sil_db(gorev_id):
    conn = sqlite3.connect(DB_YOLU)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM yapilacaklar WHERE id=?", (gorev_id,))
    conn.commit()
    conn.close()

# ---------------------------
# Yapılacaklar Penceresi
# ---------------------------
class YapilacaklarPenceresi(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Yapılacaklar Listesi")
        self.setGeometry(0, 0, 550, 520)
        self.center()
        self.secili_gorev_id = None

        layout = QVBoxLayout()
        self.setLayout(layout)
        self.setStyleSheet("background-color: #f2f2f2; color: black;")

        # Başlık label
        baslikLabel = QLabel("Yapılacaklar Listesi")
        baslikLabel.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        baslikLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(baslikLabel)

        # Başlık input
        self.baslikInput = QLineEdit()
        self.baslikInput.setPlaceholderText("Görev başlığını girin...")
        self.baslikInput.setStyleSheet("padding: 8px; border-radius: 5px; border: 1px solid #ccc; color: black;")
        layout.addWidget(self.baslikInput)

        # İçerik input
        self.detayInput = QTextEdit()
        self.detayInput.setPlaceholderText("Görev detaylarını buraya yazın...")
        self.detayInput.setStyleSheet("padding: 8px; border-radius: 5px; border: 1px solid #ccc; color: black;")
        layout.addWidget(self.detayInput)

        # Önem seçimi
        self.onemSecimi = QComboBox()
        self.onemSecimi.addItems(["Önemli (Kırmızı)", "Orta (Turuncu)", "Az Önemli (Yeşil)"])
        layout.addWidget(self.onemSecimi)

        # Görev listesi
        self.gorevList = QListWidget()
        self.gorevList.setStyleSheet("padding:5px; color: black;")
        self.gorevList.itemDoubleClicked.connect(self.gorev_duzenle)
        layout.addWidget(self.gorevList)

        # Butonlar
        btnLayout = QHBoxLayout()
        layout.addLayout(btnLayout)

        self.btnKaydet = QPushButton("Kaydet / Güncelle")
        self.btnKaydet.setStyleSheet(self.button_style())
        self.btnKaydet.clicked.connect(self.gorev_kaydet)
        btnLayout.addWidget(self.btnKaydet)

        self.btnSil = QPushButton("Seçili Görevi Sil")
        self.btnSil.setStyleSheet(self.button_style())
        self.btnSil.clicked.connect(self.gorev_sil)
        btnLayout.addWidget(self.btnSil)

        self.btnIndir = QPushButton("Görevleri İndir (.docx)")
        self.btnIndir.setStyleSheet(self.button_style())
        self.btnIndir.clicked.connect(self.gorev_indir)
        btnLayout.addWidget(self.btnIndir)

        # Veritabanından görevleri yükle
        self.gorevleri_yukle()

    # Buton stili
    def button_style(self):
        return (
            "QPushButton {background-color: #ff9800; color: white; padding: 8px; border-radius: 5px;}"
            "QPushButton:hover {background-color: #e68900;}"
        )

    # Pencereyi ortala
    def center(self):
        screen = self.screen()
        rect = screen.availableGeometry()
        x = rect.x() + (rect.width() - self.width()) // 2
        y = rect.y() + (rect.height() - self.height()) // 2
        self.move(x, y)

    # Görevleri yükle ve renk uygula
    def gorevleri_yukle(self):
        self.gorevList.clear()
        for gorev_id, baslik, _, onem in gorevleri_getir():
            item = QListWidgetItem(f"{gorev_id} - {baslik}")
            # Renk atama
            if onem == 1:
                item.setBackground(QBrush(QColor("red")))
            elif onem == 2:
                item.setBackground(QBrush(QColor("orange")))
            elif onem == 3:
                item.setBackground(QBrush(QColor("green")))
            self.gorevList.addItem(item)

    # Görevi çift tıklayarak düzenleme
    def gorev_duzenle(self, item):
        gorev_id = int(item.text().split(" - ")[0])
        conn = sqlite3.connect(DB_YOLU)
        cursor = conn.cursor()
        cursor.execute("SELECT baslik, icerik, onem FROM yapilacaklar WHERE id=?", (gorev_id,))
        baslik, icerik, onem = cursor.fetchone()
        conn.close()
        self.baslikInput.setText(baslik)
        self.detayInput.setPlainText(icerik)
        self.onemSecimi.setCurrentIndex(onem - 1)
        self.secili_gorev_id = gorev_id

    # Görev ekle / güncelle
    def gorev_kaydet(self):
        baslik = self.baslikInput.text().strip()
        icerik = self.detayInput.toPlainText().strip()
        onem = self.onemSecimi.currentIndex() + 1
        if not baslik:
            return
        if self.secili_gorev_id:
            gorev_guncelle_db(self.secili_gorev_id, baslik, icerik, onem)
        else:
            gorev_ekle_db(baslik, icerik, onem)
        self.baslikInput.clear()
        self.detayInput.clear()
        self.onemSecimi.setCurrentIndex(1)
        self.secili_gorev_id = None
        self.gorevleri_yukle()

    # Görev sil
    def gorev_sil(self):
        if self.secili_gorev_id:
            gorev_sil_db(self.secili_gorev_id)
            self.baslikInput.clear()
            self.detayInput.clear()
            self.onemSecimi.setCurrentIndex(1)
            self.secili_gorev_id = None
            self.gorevleri_yukle()

    # DOCX olarak kaydet (Masaüstü varsayılan)
    def gorev_indir(self):
        options = QFileDialog.Option.DontUseNativeDialog
        masaustu = os.path.join(os.path.expanduser("~"), "Desktop")
        dosya_yolu, _ = QFileDialog.getSaveFileName(
            self,
            "Görevleri Kaydet",
            os.path.join(masaustu, "Yapilacaklar.docx"),
            "Word Documents (*.docx)",
            options=options
        )
        if dosya_yolu:
            if not dosya_yolu.lower().endswith(".docx"):
                dosya_yolu += ".docx"
            doc = Document()
            doc.add_heading("Yapılacaklar Listesi", level=1)
            for _, baslik, icerik, onem in gorevleri_getir():
                doc.add_heading(baslik, level=2)
                doc.add_paragraph(f"[Önem Seviyesi: {onem}]")
                doc.add_paragraph(icerik)
            doc.save(dosya_yolu)

# ---------------------------
# Programı çalıştır
# ---------------------------
if __name__ == "__main__":
    veritabani_olustur()
    app = QApplication(sys.argv)
    pencere = YapilacaklarPenceresi()
    pencere.show()
    sys.exit(app.exec())
